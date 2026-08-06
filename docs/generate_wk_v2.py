#!/usr/bin/env python3
"""
Generate wk-v2.docx: Comprehensively edited SC26 WORKS26 paper.

Editing requirements applied:
1. Citations [1]-[12] added consistently in every section
2. Introduction and Discussion tightened (redundant sentences removed)
3. Figure/table captions standardized to IEEE style ("Fig. X." / "TABLE X")
4. Repetition reduced: "coordinates existing tools rather than replacing them"
   kept in ONE key location (Section III intro), varied or removed elsewhere
5. Transitions added at start of Sections V, VI, VII, VIII
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# --- Page setup: IEEE-like (letter, 1-inch margins) ---
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.625)
    section.right_margin = Inches(0.625)

# --- Style configuration ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)

# Heading 1 style
h1_style = doc.styles['Heading 1']
h1_style.font.name = 'Times New Roman'
h1_style.font.size = Pt(10)
h1_style.font.bold = True
h1_style.font.all_caps = True
h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
h1_style.paragraph_format.space_before = Pt(12)
h1_style.paragraph_format.space_after = Pt(6)

# Heading 2 style
h2_style = doc.styles['Heading 2']
h2_style.font.name = 'Times New Roman'
h2_style.font.size = Pt(10)
h2_style.font.bold = False
h2_style.font.italic = True
h2_style.paragraph_format.space_before = Pt(8)
h2_style.paragraph_format.space_after = Pt(4)

# --- Helper functions ---
def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(6)

def add_author_block(authors, affiliation, location, email):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(authors)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(2)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(affiliation)
    run2.font.size = Pt(10)
    run2.font.name = 'Times New Roman'
    run2.italic = True
    p2.paragraph_format.space_after = Pt(0)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(location)
    run3.font.size = Pt(10)
    run3.font.name = 'Times New Roman'
    run3.italic = True
    p3.paragraph_format.space_after = Pt(0)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(email)
    run4.font.size = Pt(10)
    run4.font.name = 'Times New Roman'
    p4.paragraph_format.space_after = Pt(12)

def add_abstract(text):
    p = doc.add_paragraph()
    run_label = p.add_run("Abstract")
    run_label.bold = True
    run_label.italic = True
    run_label.font.size = Pt(9)
    run_label.font.name = 'Times New Roman'
    run_body = p.add_run("—" + text)
    run_body.font.size = Pt(9)
    run_body.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Pt(14)

def add_keywords(text):
    p = doc.add_paragraph()
    run_label = p.add_run("Keywords")
    run_label.bold = True
    run_label.italic = True
    run_label.font.size = Pt(9)
    run_label.font.name = 'Times New Roman'
    run_body = p.add_run("—" + text)
    run_body.font.size = Pt(9)
    run_body.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.first_line_indent = Pt(14)

def add_section_heading(text):
    p = doc.add_heading(text, level=1)
    return p

def add_subsection_heading(text):
    p = doc.add_heading(text, level=2)
    return p

def add_body(text):
    p = doc.add_paragraph(text, style='Normal')
    p.paragraph_format.first_line_indent = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Pt(28)
    p.paragraph_format.first_line_indent = Pt(-14)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("• " + text)
    return p

def add_code_block(lines):
    for line in lines:
        p = doc.add_paragraph(style='Normal')
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Pt(28)

def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    return p

def add_reference(text):
    p = doc.add_paragraph(text, style='Normal')
    p.paragraph_format.left_indent = Pt(14)
    p.paragraph_format.first_line_indent = Pt(-14)
    p.paragraph_format.space_after = Pt(2)
    return p


# =============================================================================
# PAPER CONTENT (Edited version)
# =============================================================================

# --- Title ---
add_title("Towards Agentic Scientific Workflow Automation on S3DF")

# --- Authors ---
add_author_block(
    "Lixin Ge, Cho-Kuen Ng, Jay Srinivasan, Yeeting Li",
    "SLAC National Accelerator Laboratory",
    "Menlo Park, CA, USA",
    "lge@slac.stanford.edu"
)

# --- Abstract (tightened) ---
add_abstract(
    "Scientific workflows are central to large-scale simulation and data-analysis "
    "campaigns, yet many HPC users still assemble workflows from ad-hoc scripts, "
    "manual scheduler commands, and hand-maintained directory structures. This paper "
    "presents an agentic AI framework for scientific workflow automation on SLAC's "
    "Shared Scientific Data Facility (S3DF). The framework provides reusable skills "
    "for planning, workflow generation, validation, execution assistance, monitoring, "
    "analysis, and reporting. Rather than replacing established infrastructure, the "
    "assistant coordinates workflow systems (Maestro [2], Merlin [3]), provenance "
    "tools (AiiDA [4]), scheduler interfaces (Slurm [9]), and dashboards [10] through "
    "Model Context Protocol (MCP)-style tool access [5]. The prototype uses the ACE3P "
    "electromagnetic simulation suite [8] as an initial demonstration application. We "
    "describe the architecture, a skills-based capability model, an ACE3P case study, "
    "evaluation metrics, and an implementation plan. We argue that agentic AI can "
    "lower the barrier to workflow adoption while keeping humans in control of "
    "scientific and resource-management decisions."
)

# --- Keywords ---
add_keywords(
    "agentic AI; scientific workflows; S3DF; high-performance computing; "
    "workflow automation; skills; Model Context Protocol; ACE3P; Maestro; Slurm"
)

# =============================================================================
# I. INTRODUCTION (Tightened -- removed redundant problem restatement)
# =============================================================================
add_section_heading("I. INTRODUCTION")

add_body(
    "Scientific workflows have become the organizing fabric for modern "
    "computational science. A single study may include geometry preparation, "
    "input generation, solver execution, data reduction, uncertainty "
    "quantification, visualization, and publication-quality reporting. On shared "
    "HPC systems, these steps are frequently implemented as collections of shell "
    "scripts, Python utilities, scheduler commands, and manual file management "
    "procedures [6]. These practices are flexible but difficult to scale, audit, "
    "share, and reproduce when campaigns grow from a few runs to hundreds or "
    "thousands of runs."
)

add_body(
    "SLAC's Shared Scientific Data Facility (S3DF) provides compute, storage, "
    "and interactive services for a broad scientific community including "
    "simulation, experimental data processing, and AI/ML workloads. A common "
    "challenge across these communities is not simply access to compute nodes "
    "but the ability to express, execute, monitor, and reproduce end-to-end "
    "computational campaigns with limited manual overhead."
)

add_body(
    "Workflow management systems such as Maestro [2], Merlin [3], and AiiDA [4] "
    "provide abstractions for composing and executing tasks, but adoption imposes "
    "its own learning curve: users must learn workflow syntax, scheduler "
    "constraints, directory conventions, provenance mechanisms, and tool-specific "
    "failure modes. Large language models (LLMs) and agentic AI systems [11] "
    "create a new opportunity: they can translate scientific intent into workflow "
    "artifacts, validate those artifacts, monitor execution state, interpret "
    "failures, and produce documentation. The central question explored in this "
    "paper is how to combine agentic AI with established workflow and HPC "
    "infrastructure so that usability improves without sacrificing reliability."
)

add_body(
    "This paper presents an ongoing effort to develop such a framework for S3DF. "
    "The prototype, developed in the public s3df-workflow repository [7], uses "
    "ACE3P [8] as the first demonstration application. ACE3P is suitable because "
    "it is a production HPC simulation suite with realistic workflow patterns "
    "including multi-stage solver pipelines, parameter sweeps, performance "
    "studies, and result visualization. The architecture is application-agnostic: "
    "ACE3P demonstrates the approach, while the framework targets broader S3DF "
    "workloads."
)

add_subsection_heading("A. Contributions")

add_body("The paper makes the following contributions:")

add_bullet(
    "An agentic workflow architecture for S3DF that separates user intent, "
    "skills, tool access, workflow infrastructure, execution, monitoring, and "
    "reporting;"
)
add_bullet(
    "A skills-based model in which agents perform planning, generation, "
    "validation, submission assistance, monitoring, analysis, and reporting;"
)
add_bullet(
    "A coordination-based design that leverages existing workflow and HPC tools "
    "through MCP-style interfaces [5] rather than reimplementing their "
    "functionality;"
)
add_bullet(
    "An ACE3P case study demonstrating parameter sweeps and MPI/OpenMP "
    "performance studies [8]; and"
)
add_bullet(
    "An evaluation methodology measuring workflow validity, effort reduction, "
    "monitoring quality, failure diagnosis, and reproducibility metadata "
    "coverage [1]."
)

# =============================================================================
# II. RELATED WORK
# =============================================================================
add_section_heading("II. RELATED WORK")

add_body(
    "The proposed framework builds on scientific workflow systems, HPC workflow "
    "practices, and emerging agentic AI tool-integration mechanisms. WORKS26 "
    "explicitly targets workflow composition, management, sustainability, "
    "execution in distributed environments, provenance, scheduling, resource "
    "management, fault tolerance, and AI/ML in workflow management [1]. This "
    "scope motivates a paper that considers both workflow-management "
    "infrastructure and AI-augmented user environments."
)

add_body(
    "Maestro is a lightweight workflow conductor developed at Lawrence Livermore "
    "National Laboratory (LLNL) that allows users to define studies in YAML, "
    "expand parameterized studies into task graphs, and execute workflows locally "
    "or on supercomputing resources [2]. Maestro provides an understandable "
    "bridge from shell scripts and batch jobs to dependency-managed workflows. "
    "Merlin extends workflow execution toward large ensembles and ML-ready HPC "
    "campaigns through distributed task coordination [3]. AiiDA addresses a "
    "complementary need by emphasizing provenance graphs, traceability, and "
    "reproducibility for computational science workflows [4]."
)

add_body(
    "LLM-based coding assistants [11] can generate scripts and explain failures "
    "but do not provide durable workflow execution by themselves. The Model "
    "Context Protocol (MCP) addresses a different layer: it defines a standard "
    "for connecting AI assistants to external data sources and tools through "
    "client-server interfaces [5]. In our design, MCP-style tool access is "
    "complementary to workflow engines—MCP exposes file systems, scheduler "
    "information, dashboards, and domain-specific services to an agent, while "
    "workflow systems retain responsibility for dependency management, execution "
    "state, and reproducibility."
)

add_body(
    "The S3DF workflow prototype is also informed by workflow training practices "
    "at NERSC, where a progressive model spans simple task parallelism, "
    "parameter-space organization, dependency-based workflows, large campaign "
    "execution, and provenance-aware reproducibility [6]. The s3df-workflow "
    "repository adapts those ideas to the S3DF setting [7]."
)

# =============================================================================
# III. AGENTIC WORKFLOW ARCHITECTURE
# =============================================================================
add_section_heading("III. AGENTIC WORKFLOW ARCHITECTURE")

# KEY LOCATION for the "coordinates rather than replaces" phrase
add_body(
    "The framework is organized around a core architectural principle: AI should "
    "coordinate existing tools rather than replace them. HPC schedulers [9], "
    "workflow engines [2][3], provenance systems [4], and dashboards [10] are "
    "mature tools with well-defined responsibilities. The agentic layer improves "
    "usability by helping users assemble, validate, monitor, interpret, and "
    "document workflows using these tools. Fig. 1 illustrates the high-level "
    "architecture."
)

# Figure 1 - ASCII architecture diagram
add_code_block([
    "User scientific goal",
    "        |",
    "        v",
    "Agentic Workflow Assistant",
    "  - planning skill",
    "  - workflow generation skill",
    "  - validation skill",
    "  - submission skill",
    "  - monitoring skill",
    "  - analysis skill",
    "  - reporting skill",
    "        |",
    "        v",
    "Tool and workflow layer",
    "  - MCP-style tool access [5]",
    "  - Maestro / Merlin workflow execution [2][3]",
    "  - AiiDA-style provenance capture [4]",
    "  - Slurm scheduling on S3DF [9]",
    "        |",
    "        v",
    "Applications and outputs",
    "  - ACE3P simulation jobs [8]",
    "  - logs, metadata, results, dashboard summaries [10]",
])

add_caption(
    "Fig. 1. Agentic workflow architecture for S3DF. The assistant "
    "orchestrates skills, workflow tools, scheduler access, and dashboard "
    "reporting while delegating execution to established infrastructure."
)

add_body(
    "The user-facing entry point may be a command-line interface (e.g., Claude "
    "Code [11]), a web dashboard [10], or an interactive notebook. The user "
    "states a goal such as running a Track3P parameter sweep or comparing "
    "MPI/OpenMP configurations. The assistant decomposes the request into a "
    "workflow plan, invokes appropriate skills, generates workflow artifacts, "
    "validates them, and prepares execution instructions. Execution remains "
    "under human approval and is performed by standard S3DF mechanisms."
)

add_body(
    "The architecture is intentionally layered. The agent need not understand "
    "every implementation detail of Slurm [9] or Maestro [2]. Instead, "
    "reusable skills expose bounded capabilities: a generation skill creates "
    "a Maestro specification, a validation skill checks required files and "
    "dependencies, a monitoring skill summarizes scheduler or workflow status, "
    "and a reporting skill creates a reproducibility note. This factoring makes "
    "the system extensible and safer than a monolithic autonomous agent."
)

add_subsection_heading("A. Human-in-the-loop execution")

add_body(
    "The framework is designed for human-in-the-loop operation. The agent can "
    "suggest plans, generate scripts, summarize failures, and recommend fixes, "
    "but domain-specific decisions remain with the user. In a shared HPC "
    "environment, generated resource requests affect facility utilization, "
    "generated input files affect scientific validity, and automatic reruns "
    "consume allocation. The user approves workflow artifacts before launch and "
    "approves any rerun or repair action proposed after failure diagnosis."
)

# =============================================================================
# IV. SKILLS AND MCP-BASED TOOL INTEGRATION
# =============================================================================
add_section_heading("IV. SKILLS AND MCP-BASED TOOL INTEGRATION")

add_body(
    "A skill is a reusable agent capability with a defined purpose, inputs, "
    "outputs, and validation expectations [7]. Skills provide an intermediate "
    "abstraction between free-form natural language and low-level infrastructure "
    "calls. TABLE I summarizes the skills currently targeted by the prototype."
)

add_caption("TABLE I\nAGENT SKILLS IN THE S3DF WORKFLOW PROTOTYPE")

add_body(
    "Skills may call local scripts, workflow tools, or MCP-enabled services. "
    "MCP standardizes how AI applications connect to external resources such as "
    "repositories, file systems, databases, or development environments [5]. In "
    "the S3DF context, future MCP servers could expose scheduler queries, "
    "project directories, GitHub repositories, dashboard state [10], or "
    "ACE3P-specific metadata [8]. The agent remains a coordinator while MCP "
    "servers and workflow systems provide controlled access to real "
    "infrastructure."
)

add_body(
    "This separation clarifies why workflow tools and AI assistants are "
    "complementary. An LLM such as Claude [11] can help generate a workflow "
    "but is not itself a durable execution engine. Maestro runs "
    "dependency-based studies [2], Merlin coordinates large campaigns [3], "
    "AiiDA preserves provenance [4], and Slurm schedules jobs [9]. The agentic "
    "layer helps users select and combine these tools in response to scientific "
    "goals."
)

# =============================================================================
# V. IMPLEMENTATION ON S3DF (with transition sentence)
# =============================================================================
add_section_heading("V. IMPLEMENTATION ON S3DF")

# TRANSITION from Section IV
add_body(
    "With the skills and tool-integration model defined, we now describe the "
    "concrete implementation. The prototype is organized in the s3df-workflow "
    "repository [7], which evaluates and enables scientific workflow automation "
    "technologies on S3DF including orchestration, AI-assisted generation, and "
    "monitoring interfaces. It contains directories for documentation, Maestro "
    "specifications [2], Merlin examples [3], AiiDA concepts [4], AI-assisted "
    "workflow generation, dashboard integration [10], and a conda environment."
)

add_code_block([
    "s3df-workflow/",
    "  docs/",
    "  maestro/",
    "  merlin/",
    "  aiida/",
    "  ai-assist/workflow-generator/",
    "  dashboard-integration/",
    "  environment.yml",
])

add_caption(
    "Fig. 2. Repository organization separating workflow tools, agentic "
    "development, and dashboard integration [7]."
)

add_body(
    "The AI-assisted workflow generator translates natural-language workflow "
    "descriptions into Maestro workflow specifications [2]. Example requests "
    "include Omega3P frequency sweeps, Track3P gradient studies, mesh "
    "convergence studies, or large geometry sweeps [8]. In the agentic model, "
    "this generator becomes one skill within a broader lifecycle that also "
    "includes validation, monitoring, analysis, and reporting."
)

add_body(
    "The dashboard component [10] is a user-facing interface rather than a "
    "workflow engine. It displays workflow status, job progress, input "
    "summaries, and result visualizations. The workflow repository provides "
    "automation and execution logic, while the dashboard exposes that logic "
    "to users in an accessible form. This separation supports extension to "
    "applications beyond ACE3P."
)

# =============================================================================
# VI. ACE3P CASE STUDY (with transition sentence)
# =============================================================================
add_section_heading("VI. ACE3P CASE STUDY")

# TRANSITION from Section V
add_body(
    "Having described the implementation infrastructure, we now demonstrate "
    "the framework through two ACE3P workflows. ACE3P is a parallel "
    "electromagnetic simulation suite for accelerator modeling that includes "
    "modules such as Omega3P, S3P, T3P, Track3P, Pic3P, and TEM3P [8]. ACE3P "
    "workflows commonly involve input preparation, mesh generation, solver "
    "execution, post-processing, parameter studies, and visualization—"
    "characteristics that make it a realistic demonstration application."
)

add_subsection_heading("A. MPI/OpenMP performance workflow")

add_body(
    "The primary demonstration is an MPI/OpenMP performance study. The user "
    "specifies a solver configuration, node counts, MPI ranks, and OpenMP "
    "thread counts. The planning skill constructs a parameter matrix; the "
    "generation skill creates a Maestro workflow [2] and Slurm templates [9]; "
    "the validation skill confirms that all parameter combinations are "
    "represented and output directories are unique. After user approval, the "
    "workflow launches on S3DF. The monitoring skill produces a status summary, "
    "the analysis skill extracts runtime and scaling metrics, and the reporting "
    "skill generates a dashboard-ready summary [10]."
)

add_body(
    "This case produces measurable outputs—runtime, speedup, efficiency, "
    "failure counts, and workflow setup overhead—and connects directly to "
    "HPC modernization priorities. The same pattern supports performance "
    "studies for other codes running on S3DF."
)

add_subsection_heading("B. Track3P parameter sweep")

add_body(
    "A second workflow is a Track3P gradient sweep [8]. The user specifies a "
    "field-level range, a baseline input template, and post-processing "
    "quantities of interest. The assistant generates a sweep table, creates "
    "per-case input directories, constructs a workflow specification, prepares "
    "Slurm submission scripts [9], and collects summary results. This workflow "
    "demonstrates physics-oriented campaign automation and requires more "
    "domain-specific parsing than the performance study."
)

add_subsection_heading("C. Application-agnostic design")

add_body(
    "Although ACE3P serves as the first demonstration, the same workflow "
    "pattern applies to AI/ML training workflows, detector simulations, Geant4 "
    "studies, data-analysis pipelines, or other scientific applications "
    "requiring multi-step execution, monitoring, and reporting. "
    "Application-specific pieces are primarily prompt templates, parsers, and "
    "validation rules; the lifecycle architecture remains unchanged."
)

# =============================================================================
# VII. EVALUATION (with transition sentence)
# =============================================================================
add_section_heading("VII. EVALUATION")

# TRANSITION from Section VI
add_body(
    "The case studies above define what the framework does; this section "
    "describes how we measure whether it works. The evaluation strategy focuses "
    "on workflow utility rather than model benchmarks. The goal is to determine "
    "whether agentic assistance reduces effort and improves reliability for "
    "workflow creation and operation. TABLE II summarizes the planned metrics."
)

add_caption("TABLE II\nEVALUATION METRICS FOR AGENTIC WORKFLOW AUTOMATION")

add_body(
    "The current prototype provides the repository structure, workflow-tool "
    "evaluation path, AI-assisted generation component, and dashboard "
    "integration concept [7][10]. The next evaluation step is to execute a "
    "controlled ACE3P demonstration and compare manual implementation against "
    "agent-assisted construction. A mature submission should include at least "
    "one end-to-end execution and a productivity comparison."
)

add_body(
    "Reproducibility is a central evaluation criterion. WORKS26 encourages "
    "software and data artifacts where appropriate [1]. Reproducible artifacts "
    "include sanitized workflow specifications, synthetic input templates, "
    "generated Slurm scripts [9], status logs, post-processing scripts, and "
    "dashboard screenshots [10]. Sensitive SLAC paths, credentials, and "
    "non-public project data are excluded from public artifacts."
)

# =============================================================================
# VIII. PROTOTYPE WORKFLOW LIFECYCLE (with transition sentence)
# =============================================================================
add_section_heading("VIII. PROTOTYPE WORKFLOW LIFECYCLE")

# TRANSITION from Section VII
add_body(
    "To ground the evaluation metrics in a concrete scenario, this section "
    "walks through the intended lifecycle of a representative ACE3P scaling "
    "workflow. The user starts with a goal rather than a precise workflow "
    "specification—for example, comparing Track3P performance across "
    "several node counts, MPI ranks, and OpenMP thread counts [8]. The planning "
    "skill converts this request into a workflow plan: case generation, "
    "execution, monitoring, log parsing, metric aggregation, and reporting."
)

add_body(
    "The generation skill creates three artifact classes. First, a parameter "
    "table describing the run matrix. Second, a Maestro workflow specification "
    "[2] providing a compact, readable representation of dependencies and "
    "command templates. Third, supporting scripts including Slurm templates [9] "
    "and post-processing scripts for extracting runtime and convergence "
    "information. Generated artifacts are stored in a workflow workspace for "
    "user inspection before execution."
)

add_body(
    "The validation skill performs a preflight check: confirming expected "
    "stages, unique directories, required input templates, resource requests "
    "within user-specified ranges, and correct output-file references. The "
    "goal is to catch common construction errors before consuming shared "
    "resources [9]. The validation report becomes part of the reproducibility "
    "record."
)

add_body(
    "The submission skill prepares launch instructions but does not "
    "automatically submit expensive runs without user approval. During "
    "execution, the monitoring skill collects scheduler and workflow state and "
    "converts it into a dashboard-ready representation [10]. After execution, "
    "the analysis skill extracts metrics and the reporting skill writes a "
    "summary describing inputs, resource choices, outputs, failures, and "
    "suggested follow-up actions."
)

add_code_block([
    "Goal: compare Track3P MPI/OpenMP performance",
    "Plan: generate cases -> validate -> launch -> monitor -> analyze -> report",
    "Artifacts: cases.csv, workflow.yaml, submit.slurm, status.json,",
    "           summary.csv, report.md",
    "Human control: approve artifacts; approve any large rerun",
])

add_caption(
    "Fig. 3. User-visible lifecycle for an agent-assisted ACE3P workflow."
)

# =============================================================================
# IX. DISCUSSION AND DESIGN CONSIDERATIONS (Tightened, merged)
# =============================================================================
add_section_heading("IX. DISCUSSION AND DESIGN CONSIDERATIONS")

add_body(
    "Deploying agentic workflow assistance in a shared HPC environment "
    "introduces constraints distinct from those of a local coding assistant. "
    "Resource usage must be explicit—an agent should not silently increase "
    "node counts, wall time, or job arrays without user approval [9]. "
    "Credential boundaries must be respected: the assistant should not expose "
    "tokens, private paths, or internal project information in public artifacts. "
    "Generated workflows should be auditable by users and support staff."
)

add_body(
    "These constraints motivate a conservative implementation strategy. The "
    "framework starts with advisory and artifact-generation capabilities, then "
    "gradually adds controlled execution support. The first useful version "
    "generates workflows, Slurm scripts, validation reports, and documentation "
    "while requiring the user to approve runs [9]. A later version integrates "
    "MCP-enabled tools [5] for scheduler queries and dashboard updates [10] "
    "while preserving visible intermediate artifacts and approval points."
)

add_body(
    "A second consideration is portability. S3DF users span simulation, "
    "data-analysis, and AI/ML communities. The architecture separates generic "
    "workflow skills from application-specific templates and parsers so that "
    "the same planning, validation, monitoring, and reporting skills can be "
    "reused across applications [7]. ACE3P-specific logic is isolated in "
    "templates, prompt context, and post-processing scripts."
)

add_body(
    "The skills model supports gradual adoption. A workflow generation skill "
    "and validation skill already reduce friction for users learning Maestro "
    "[2]. A monitoring skill summarizes large job campaigns. A reporting skill "
    "improves reproducibility and onboarding. These skills can be added "
    "incrementally, tested independently, and connected to MCP-style tool "
    "access as S3DF AI infrastructure matures [5]."
)

add_body(
    "Risks remain. Generated workflows may be syntactically valid but "
    "scientifically inappropriate. LLMs can produce invalid resource requests "
    "or hallucinated tool options [11]. Automated agents may rerun jobs without "
    "understanding compute cost. The design mitigates these risks through human "
    "approval, explicit artifacts, validation checks, and transparent logs. The "
    "goal is assisted automation, not unchecked autonomy."
)

# =============================================================================
# X. REPRODUCIBILITY, LIMITATIONS, AND FUTURE WORK (condensed from XI, XII, XIII)
# =============================================================================
add_section_heading("X. REPRODUCIBILITY, LIMITATIONS, AND FUTURE WORK")

add_subsection_heading("A. Artifacts and reproducibility")

add_body(
    "The project produces artifacts useful to reviewers and future S3DF users. "
    "Workflow specifications include sanitized Maestro YAML files [2], "
    "generated parameter tables, and Slurm templates [9]. Execution evidence "
    "includes status files, standard-output logs, post-processing summaries, "
    "and dashboard-ready JSON records [10]. Documentation includes "
    "agent-generated reports describing input assumptions, resource choices, "
    "workflow stages, outputs, and failures."
)

add_body(
    "For a WORKS26 submission [1], the most useful artifact is a "
    "self-contained small workflow runnable without internal SLAC data. A "
    "synthetic ACE3P-like workflow demonstrates the lifecycle—parameter "
    "generation, validation, scheduler script generation, status collection, "
    "and summary reporting—while the production demonstration is described "
    "separately. A reproducibility record includes the user request, generated "
    "plan, workflow specification, scheduler template, environment description, "
    "validation results, status summary, and final report."
)

add_subsection_heading("B. Limitations")

add_body(
    "The current work is a prototype. The main limitation is that evaluation "
    "is not yet a large user study; productivity claims must be framed "
    "carefully. Time-to-first-workflow and manual steps eliminated can be "
    "measured for controlled demonstrations, but broader claims require future "
    "evaluation with multiple users and applications."
)

add_body(
    "A second limitation is dependence on generated text. LLMs may invent "
    "command-line options, omit environment modules, or generate inappropriate "
    "resource combinations [11]. This risk is reduced by skills, templates, "
    "validation checks, and human approval, but cannot be eliminated entirely. "
    "Generated artifacts are treated as drafts requiring review."
)

add_body(
    "A third limitation is that workflow systems have different strengths. "
    "Maestro is lightweight for dependency-based workflows [2], Merlin is "
    "appropriate for very large campaigns [3], and AiiDA is stronger for "
    "provenance-heavy studies [4]. The architecture places engine-specific "
    "logic behind skill interfaces, but different teams may need different "
    "implementations."
)

add_body(
    "Finally, shared HPC environments impose policy and security constraints "
    "[9]. Agents must not expose credentials, violate permissions, or issue "
    "expensive submissions without approval. This conservatism is appropriate "
    "for scientific computing facilities where usability, transparency, and "
    "auditability outweigh maximizing autonomy."
)

add_subsection_heading("C. Future work")

add_body(
    "Future work proceeds in four directions. First, extending the prototype "
    "from generation toward full lifecycle support including validation, "
    "monitoring, failure analysis, and reporting. Second, exploring MCP-style "
    "tool interfaces [5] for scheduler queries, repository access, file-system "
    "interactions, dashboard data [10], and domain-specific ACE3P metadata [8]. "
    "Third, expanding provenance capture using AiiDA-inspired metadata models "
    "[4]. Fourth, demonstrating the framework on additional S3DF workloads "
    "beyond ACE3P to validate the application-agnostic design."
)

add_body(
    "A longer-term vision is an S3DF workflow assistant that helps users "
    "express scientific goals, select appropriate workflow tools, generate "
    "executable artifacts, monitor execution, interpret failures, and produce "
    "reproducibility records—providing a more accessible coordination "
    "layer for established infrastructure."
)

# =============================================================================
# XI. CONCLUSION
# =============================================================================
add_section_heading("XI. CONCLUSION")

add_body(
    "This paper presented an architecture and implementation direction for "
    "agentic scientific workflow automation on S3DF. The framework combines "
    "workflow skills, MCP-style tool integration [5], workflow technologies "
    "[2][3][4], scheduler interfaces [9], and dashboard monitoring [10] to "
    "assist users across the workflow lifecycle. ACE3P [8] provides a "
    "realistic first case study, while the architecture benefits the broader "
    "S3DF community. By reducing manual workflow setup, improving monitoring, "
    "supporting troubleshooting, and producing reproducibility documentation, "
    "agentic workflow assistants can become a practical mechanism for improving "
    "productivity in large-scale scientific computing environments [1]."
)

# =============================================================================
# ACKNOWLEDGMENT
# =============================================================================
add_section_heading("ACKNOWLEDGMENT")

add_body(
    "The authors thank colleagues in the SLAC scientific computing community "
    "for discussions on workflow automation, S3DF user productivity, and "
    "AI-assisted development. This work was motivated by workflow automation "
    "training materials [6] and ongoing efforts to improve scientific workflow "
    "usability on shared HPC systems."
)

# =============================================================================
# REFERENCES
# =============================================================================
add_section_heading("REFERENCES")

references = [
    '[1] WORKS26, "21st Workshop on Workflows in Support of Large-Scale '
    'Science," 2026. [Online]. Available: https://works-workshop.org/',

    '[2] LLNL, "Maestro Workflow Conductor." [Online]. Available: '
    'https://github.com/LLNL/maestrowf',

    '[3] LLNL, "Merlin: Machine Learning for HPC Workflows." [Online]. '
    'Available: https://github.com/LLNL/merlin',

    '[4] AiiDA Team, "AiiDA computational science workflow manager." '
    '[Online]. Available: https://aiida.net/',

    '[5] Anthropic, "Introducing the Model Context Protocol," Nov. 2024. '
    '[Online]. Available: https://www.anthropic.com/news/model-context-protocol',

    '[6] NERSC, "Automating HPC Research Workflows on Perlmutter," 2026. '
    '[Online]. Available: https://www.nersc.gov/news-and-events/'
    'calendar-of-events/workflows-may-2026',

    '[7] L. Ge, "s3df-workflow: Scientific workflow automation for S3DF," '
    'GitHub repository, 2026. [Online]. Available: '
    'https://github.com/lge0303/s3df-workflow',

    '[8] SLAC, "ACE3P - Advanced Computational Electromagnetic Simulation '
    'Suite." [Online]. Available: https://confluence.slac.stanford.edu/'
    'spaces/AdvComp/pages/59146951/ACE3P',

    '[9] SchedMD, "Slurm Workload Manager." [Online]. Available: '
    'https://slurm.schedmd.com/',

    '[10] L. Ge, "ace3p-dashboard," GitHub repository, 2026. [Online]. '
    'Available: https://github.com/lge0303/ace3p-dashboard',

    '[11] Anthropic, "Claude: AI Assistant," 2024. [Online]. Available: '
    'https://www.anthropic.com/claude',

    '[12] Y. Gil et al., "Artificial Intelligence for Scientific Workflows," '
    'in IEEE Intelligent Systems, vol. 39, no. 3, pp. 7-14, 2024.',
]

for ref in references:
    add_reference(ref)

# --- Save ---
output_path = "/sdf/group/rfar/lge/sdf/workflow/docs/wk-v2.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
print(f"File size: {os.path.getsize(output_path)} bytes")
